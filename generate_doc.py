from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    doc = Document()
    doc.add_heading("Animal Detection V2 Documentation", level=1)

    doc.add_heading("Project Overview", level=2)
    doc.add_paragraph(
        "Detect six animal species (buffaloes, deers, elephants, rhinos, tigers, wild boars) "
        "from camera feeds using a YOLO model. This project provides scripts to split the dataset, "
        "train the model, and run live inference with audio alerts."
    )

    doc.add_heading("Repository Structure", level=2)
    structure_items = [
        "Animal Detection V2/",
        "images/ – Source images (flat structure)",
        "labels/ – YOLO annotation files",
        "classes.txt – Class names",
        "notes.json – Dataset metadata",
        "datasets/animals/ – Split dataset (train/val/test)",
        "runs/animals/ – Ultralytics training outputs",
        "animals.yaml – Dataset configuration for Ultralytics",
        "prepare_dataset.py – Dataset splitting utility",
        "train_yolo.py – Training script wrapper",
        "detect_and_alert.py – Live detection with audio alert",
        "requirements.txt – Python dependencies",
    ]
    for item in structure_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Setup", level=2)
    doc.add_paragraph("1. Create and activate a virtual environment:", style="List Number")
    doc.add_paragraph("python -m venv .venv", style="List Continue")
    doc.add_paragraph(".\\.venv\\Scripts\\activate", style="List Continue")
    doc.add_paragraph("pip install --upgrade pip", style="List Continue")
    doc.add_paragraph("pip install -r requirements.txt", style="List Continue")
    doc.add_paragraph("2. Split the dataset into train/val/test:", style="List Number")
    doc.add_paragraph("python prepare_dataset.py", style="List Continue")
    doc.add_paragraph("Optional flags: --ratios, --seed, --move", style="List Continue")

    doc.add_heading("Training", level=2)
    doc.add_paragraph("python train_yolo.py --model yolov8n.pt --device 0 --epochs 100")
    doc.add_paragraph("Key options:", style="List Bullet")
    doc.add_paragraph("--batch: batch size", style="List Continue")
    doc.add_paragraph("--imgsz: input image size", style="List Continue")
    doc.add_paragraph("--project / --name: run directory naming", style="List Continue")
    doc.add_paragraph("--resume: continue previous training", style="List Continue")
    doc.add_paragraph(
        "Weights are exported to runs/animals/<run-name>/weights/{best,last}.pt."
    )

    doc.add_heading("Live Detection", level=2)
    doc.add_paragraph("python detect_and_alert.py --source 0 --audio path\\to\\alert.wav")
    doc.add_paragraph("--model: override weights path", style="List Bullet")
    doc.add_paragraph("--conf: detection confidence threshold", style="List Continue")
    doc.add_paragraph("--cooldown: seconds between audio alerts", style="List Continue")
    doc.add_paragraph("--device: inference device (GPU id or cpu)", style="List Continue")
    doc.add_paragraph(
        "Automatically falls back to the most recent weights in runs/animals/**/weights/."
    )

    doc.add_heading("Dependencies", level=2)
    doc.add_paragraph("Installed through requirements.txt:", style="List Bullet")
    doc.add_paragraph("ultralytics – Training/inference engine", style="List Continue")
    doc.add_paragraph("opencv-python – Video capture/visualisation", style="List Continue")
    doc.add_paragraph("playsound – Audio alerts", style="List Continue")
    doc.add_paragraph("polars – Ultralytics CSV logging backend", style="List Continue")
    doc.add_paragraph("python-docx – Documentation generation (optional)", style="List Continue")

    doc.add_heading("Tips", level=2)
    doc.add_paragraph("Monitor training metrics under runs/animals/<run-name>/", style="List Bullet")
    doc.add_paragraph("Test on recorded videos via --source path\\to\\video.mp4", style="List Continue")
    doc.add_paragraph("Ensure GPU drivers/CUDA are configured for accelerated training", style="List Continue")

    output_path = Path("Animal Detection V2 Documentation.docx")
    doc.save(output_path)
    print(f"Wrote documentation to {output_path.resolve()}")


if __name__ == "__main__":
    main()
